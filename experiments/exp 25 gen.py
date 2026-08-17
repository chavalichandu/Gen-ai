"""
Document-Based AI Assistant: Multiple Uploaded Documents + Q&A
--------------------------------------------------------------------
Supports uploading multiple documents in different formats (.txt,
.pdf, .docx), manages them as a named collection, and answers
questions by finding and citing the most relevant uploaded document(s).

This is the "simple assistant" from item 9, extended with:
  - Multi-format file loading (.txt, .pdf, .docx) -- real "uploads",
    not just plain text
  - A persistent document collection: add / list / remove by name
  - Retrieval across ALL uploaded documents with per-document citation

For long documents that should be split into passages instead of
matched whole, see rag_pipeline.py's chunking -- this version treats
each uploaded file as one retrievable unit, which fits typical "upload
a few files and ask questions" use.

Install:
    pip install scikit-learn numpy pypdf python-docx
    pip install anthropic     # optional, for real generated answers
"""

from __future__ import annotations
import os
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer


# ----------------------------------------------------------------------
# File loading: extract text depending on format
# ----------------------------------------------------------------------
def extract_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".txt":
        with open(filepath, encoding="utf-8") as f:
            return f.read()

    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if ext == ".docx":
        import docx
        doc = docx.Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)

    raise ValueError(f"Unsupported file type: {ext} (supported: .txt, .pdf, .docx)")


# ----------------------------------------------------------------------
# Document-based assistant
# ----------------------------------------------------------------------
class DocumentAssistant:
    def __init__(self):
        self.filenames: list[str] = []
        self.texts: list[str] = []
        self.vectorizer: TfidfVectorizer | None = None
        self.doc_vectors = None

    def upload(self, filepath: str):
        """Add a single uploaded document (.txt, .pdf, or .docx) to the collection."""
        text = extract_text(filepath)
        if not text.strip():
            raise ValueError(f"No extractable text found in '{filepath}' "
                              f"(scanned PDFs need OCR first).")
        self.filenames.append(os.path.basename(filepath))
        self.texts.append(text)
        self._reindex()

    def upload_many(self, filepaths: list[str]):
        for path in filepaths:
            self.upload(path)

    def remove(self, filename: str):
        """Remove a document from the collection by its filename."""
        if filename not in self.filenames:
            raise KeyError(f"'{filename}' is not in the collection.")
        i = self.filenames.index(filename)
        del self.filenames[i]
        del self.texts[i]
        self._reindex()

    def list_documents(self) -> list[str]:
        return list(self.filenames)

    def _reindex(self):
        """Rebuild the TF-IDF index over all currently uploaded documents."""
        if not self.texts:
            self.vectorizer, self.doc_vectors = None, None
            return
        self.vectorizer = TfidfVectorizer(stop_words="english")
        self.doc_vectors = self.vectorizer.fit_transform(self.texts)

    def top_matches(self, question: str, top_k: int = 3) -> list[dict]:
        """Rank all uploaded documents by relevance to the question."""
        if not self.texts:
            return []
        question_vector = self.vectorizer.transform([question])
        scores = (self.doc_vectors @ question_vector.T).toarray().flatten()
        top_k = min(top_k, len(self.texts))
        top_indices = np.argsort(scores)[::-1][:top_k]
        return [
            {"filename": self.filenames[i], "text": self.texts[i], "score": float(scores[i])}
            for i in top_indices
        ]

    def ask(self, question: str, top_k: int = 3, use_llm: bool = True) -> dict:
        """Answer a question using the most relevant uploaded document(s), with citations."""
        matches = self.top_matches(question, top_k)
        matches = [m for m in matches if m["score"] > 0]

        if not matches:
            return {"answer": "None of the uploaded documents seem relevant to that question.",
                     "sources": []}

        if use_llm and os.environ.get("ANTHROPIC_API_KEY"):
            answer = self._generate_with_llm(question, matches)
        else:
            answer = (f"(no LLM configured -- most relevant excerpt, from "
                      f"'{matches[0]['filename']}':)\n{matches[0]['text'][:500]}")

        return {
            "answer": answer,
            "sources": [{"filename": m["filename"], "score": round(m["score"], 4)} for m in matches],
        }

    def _generate_with_llm(self, question: str, matches: list[dict]) -> str:
        import anthropic
        client = anthropic.Anthropic()
        context = "\n\n".join(f"[{m['filename']}]\n{m['text']}" for m in matches)
        prompt = (f"Answer the question using only the documents below. Cite which "
                  f"document(s) you used by filename. If the answer isn't there, say so.\n\n"
                  f"{context}\n\nQuestion: {question}")
        response = client.messages.create(
            model="claude-sonnet-4-6", max_tokens=400,
            messages=[{"role": "user", "content": prompt}],
        )
        return "".join(b.text for b in response.content if b.type == "text")


# ----------------------------------------------------------------------
# Demo: upload a .txt, a .pdf, and a .docx, then ask questions across all three
# ----------------------------------------------------------------------
if __name__ == "__main__":
    os.makedirs("/tmp/uploads", exist_ok=True)

    # .txt upload
    with open("/tmp/uploads/company_policy.txt", "w") as f:
        f.write("Employees are entitled to 20 days of paid vacation per year, "
                 "accrued monthly. Unused days roll over up to a maximum of 5 days.")

    # .pdf upload (generated for the demo)
    from pypdf import PdfWriter
    from reportlab.pdfgen import canvas
    c = canvas.Canvas("/tmp/uploads/product_spec.pdf")
    c.drawString(100, 750, "The X200 model has a battery life of 18 hours and")
    c.drawString(100, 730, "supports fast charging to 80% in 30 minutes.")
    c.save()

    # .docx upload
    import docx
    d = docx.Document()
    d.add_paragraph("The onboarding process takes approximately 2 weeks and includes "
                     "IT setup, HR paperwork, and a manager introduction meeting.")
    d.save("/tmp/uploads/onboarding_guide.docx")

    assistant = DocumentAssistant()
    assistant.upload_many([
        "/tmp/uploads/company_policy.txt",
        "/tmp/uploads/product_spec.pdf",
        "/tmp/uploads/onboarding_guide.docx",
    ])
    print("Uploaded documents:", assistant.list_documents())

    for question in [
        "How many vacation days do employees get?",
        "What's the battery life of the X200?",
        "How long does onboarding take?",
    ]:
        result = assistant.ask(question)
        print(f"\nQ: {question}")
        print(f"A: {result['answer']}")
        print(f"Sources: {result['sources']}")
